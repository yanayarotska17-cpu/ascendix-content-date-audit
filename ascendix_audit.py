import requests
import xml.etree.ElementTree as ET
import time
import sys
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SITEMAP_INDEX_URL = "https://ascendix.com/sitemap_index.xml"
NS = {"sm": "http://www.sitemaps.org/schemas/sitemap/0.9"}
DELAY = 1


def fetch_xml(url):
    try:
        resp = requests.get(url, timeout=30, headers={"User-Agent": "Mozilla/5.0 (audit-bot)"})
        resp.raise_for_status()
        return ET.fromstring(resp.content)
    except requests.HTTPError as e:
        print(f"HTTP error fetching {url}: {e}", file=sys.stderr)
        return None
    except Exception as e:
        print(f"Error fetching {url}: {e}", file=sys.stderr)
        return None


def get_child_sitemap_urls(index_root):
    urls = []
    for loc in index_root.findall(".//sm:sitemap/sm:loc", NS):
        urls.append(loc.text.strip())
    return urls


def parse_urlset(root):
    entries = []
    for url_el in root.findall("sm:url", NS):
        loc = url_el.find("sm:loc", NS)
        lastmod = url_el.find("sm:lastmod", NS)
        if loc is None:
            continue
        url = loc.text.strip()
        mod = lastmod.text.strip() if lastmod is not None else None
        entries.append((url, mod))
    return entries


def classify(url):
    if "/blog/" in url:
        return "blog"
    if "/cases/" in url:
        return "cases"
    return "pages"


def sort_key(entry):
    _, mod = entry
    if mod is None:
        return "9999-99-99"
    return mod


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_section_table(doc, title, entries):
    heading = doc.add_heading(level=2)
    run = heading.add_run(f"{title}  ({len(entries)} pages)")
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x9E)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(1.8)

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["URL", "Last Modified Date"]):
        set_cell_bg(cell, "2C5F9E")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for i, (url, mod) in enumerate(entries):
        row = table.add_row().cells
        bg = "F4F7FC" if i % 2 == 0 else "FFFFFF"
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
        p0 = row[0].paragraphs[0]
        p0.add_run(url).font.size = Pt(9)
        p1 = row[1].paragraphs[0]
        p1.add_run(mod if mod else "unknown").font.size = Pt(9)

    doc.add_paragraph()


def main():
    print("Fetching sitemap index...")
    index_root = fetch_xml(SITEMAP_INDEX_URL)
    if index_root is None:
        print("Failed to fetch sitemap index. Exiting.", file=sys.stderr)
        sys.exit(1)

    child_urls = get_child_sitemap_urls(index_root)
    print(f"Found {len(child_urls)} child sitemaps.")

    all_entries = []
    for child_url in child_urls:
        print(f"  Fetching: {child_url}")
        time.sleep(DELAY)
        child_root = fetch_xml(child_url)
        if child_root is None:
            continue
        entries = parse_urlset(child_root)
        print(f"    -> {len(entries)} URLs")
        all_entries.extend(entries)

    filtered = [(url, mod) for url, mod in all_entries
                if "category" not in url and "author" not in url]
    print(f"Total URLs after filtering: {len(filtered)}")

    blog, cases, pages = [], [], []
    for entry in filtered:
        t = classify(entry[0])
        if t == "blog":
            blog.append(entry)
        elif t == "cases":
            cases.append(entry)
        else:
            pages.append(entry)

    blog.sort(key=sort_key)
    cases.sort(key=sort_key)
    pages.sort(key=sort_key)

    today = date.today().isoformat()
    filename = f"ascendix_audit_{today}.docx"

    doc = Document()

    title = doc.add_heading(level=1)
    run = title.add_run(f"Ascendix Site Audit — {today}")
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    summary = doc.add_paragraph()
    summary.add_run(
        f"Blog posts: {len(blog)}    |    Cases: {len(cases)}    |    Pages: {len(pages)}    |    Total: {len(blog)+len(cases)+len(pages)}"
    ).font.size = Pt(11)
    doc.add_paragraph()

    add_section_table(doc, "Blog Posts", blog)
    add_section_table(doc, "Cases", cases)
    add_section_table(doc, "Pages", pages)

    doc.save(filename)
    print(f"Report saved: {filename}")


if __name__ == "__main__":
    main()
